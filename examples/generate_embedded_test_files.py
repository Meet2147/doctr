from __future__ import annotations

import os
import zipfile
from pathlib import Path


def _ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def create_embedded_text_files(out_dir: Path) -> tuple[Path, Path]:
    _ensure_dir(out_dir)
    txt = out_dir / "embedded_note.txt"
    md = out_dir / "embedded_child.md"
    txt.write_text("This is an embedded TXT file for doctr testing.", encoding="utf-8")
    md.write_text("# Embedded Child\n\nThis markdown file is embedded inside a container.", encoding="utf-8")
    return txt, md


def create_docx_with_embedded_files(out_dir: Path, txt: Path, md: Path) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install office deps first: pip install 'doctr-index[office]'") from exc

    _ensure_dir(out_dir)
    docx_path = out_dir / "host_with_embedded.docx"

    doc = Document()
    doc.add_heading("Host DOCX", level=1)
    doc.add_paragraph("This DOCX has embedded files under word/embeddings/.")
    doc.save(docx_path)

    with zipfile.ZipFile(docx_path, "a", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("word/embeddings/embedded_note.txt", txt.read_bytes())
        zf.writestr("word/embeddings/embedded_child.md", md.read_bytes())

    return docx_path


def create_xlsx_with_embedded_files(out_dir: Path, txt: Path, md: Path) -> Path:
    try:
        from openpyxl import Workbook
    except ImportError as exc:
        raise RuntimeError("Install office deps first: pip install 'doctr-index[office]'") from exc

    _ensure_dir(out_dir)
    xlsx_path = out_dir / "host_with_embedded.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Overview"
    ws["A1"] = "Host XLSX"
    ws["A2"] = "This workbook has embedded files under xl/embeddings/."
    wb.save(xlsx_path)

    with zipfile.ZipFile(xlsx_path, "a", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("xl/embeddings/embedded_note.txt", txt.read_bytes())
        zf.writestr("xl/embeddings/embedded_child.md", md.read_bytes())

    return xlsx_path


def main() -> None:
    base = Path("examples")
    embedded_dir = base / "embedded"
    hosts_dir = base / "hosts"
    _ensure_dir(base)

    txt, md = create_embedded_text_files(embedded_dir)
    docx_path = create_docx_with_embedded_files(hosts_dir, txt, md)
    xlsx_path = create_xlsx_with_embedded_files(hosts_dir, txt, md)

    print("Created test files:")
    print(f"- {txt}")
    print(f"- {md}")
    print(f"- {docx_path}")
    print(f"- {xlsx_path}")
    print("\nTry:")
    print("  from doctr import index_document")
    print(f"  idx = index_document('{docx_path}', include_embedded=True)")
    print("  print(idx.to_pageindex_dict(include_empty_nodes=False))")


if __name__ == "__main__":
    main()
