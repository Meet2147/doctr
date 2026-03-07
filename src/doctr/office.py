from __future__ import annotations

import os
import re
from dataclasses import dataclass
from datetime import datetime

from .models import DocumentIndex, SectionNode


@dataclass(slots=True)
class HeadingEntry:
    title: str
    level: int
    start_para: int


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def index_docx_file(path: str) -> DocumentIndex:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Missing dependency `python-docx`. Install with: pip install 'doctr[office]'") from exc

    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    styles = [getattr(getattr(p, "style", None), "name", "") or "" for p in doc.paragraphs]

    headings: list[HeadingEntry] = []
    for i, (txt, style_name) in enumerate(zip(paragraphs, styles), start=1):
        if not txt:
            continue
        style = style_name.lower()
        if style.startswith("heading"):
            m = re.search(r"(\d+)", style)
            level = int(m.group(1)) if m else 1
            headings.append(HeadingEntry(title=_normalize(txt), level=max(1, level), start_para=i))

    def summarize(start: int, end: int) -> str:
        chunk = [_normalize(p) for p in paragraphs[start - 1 : end] if _normalize(p)]
        if not chunk:
            return "[No extractable text in this section]"
        return " ".join(chunk)

    if not headings:
        full_text = " ".join(_normalize(p) for p in paragraphs if _normalize(p))
        nodes = [
            SectionNode(
                title=os.path.basename(path),
                start_page=1,
                end_page=max(1, len(paragraphs)),
                summary=full_text or "[No extractable text in this document]",
            )
        ]
    else:
        roots: list[SectionNode] = []
        stack: list[tuple[int, SectionNode]] = []
        flat: list[SectionNode] = []

        for h in headings:
            node = SectionNode(title=h.title, start_page=h.start_para)
            while stack and stack[-1][0] >= h.level:
                stack.pop()
            if stack:
                stack[-1][1].children.append(node)
            else:
                roots.append(node)
            stack.append((h.level, node))
            flat.append(node)

        for i, n in enumerate(flat):
            next_start = flat[i + 1].start_page if i + 1 < len(flat) else len(paragraphs) + 1
            n.end_page = max(n.start_page or 1, next_start - 1)
            n.summary = summarize(n.start_page or 1, n.end_page or (n.start_page or 1))

        nodes = roots

    stat = os.stat(path)
    metadata = {
        "file_name": os.path.basename(path),
        "file_path": path,
        "file_size_bytes": stat.st_size,
        "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "extension": ".docx",
        "indexing": {"index_source": "docx_headings", "unit": "paragraph"},
    }

    return DocumentIndex(
        title=nodes[0].title if nodes else None,
        total_pages=len(paragraphs) if paragraphs else 0,
        nodes=nodes,
        source=path,
        metadata=metadata,
    )


def index_xlsx_file(path: str) -> DocumentIndex:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("Missing dependency `openpyxl`. Install with: pip install 'doctr[office]'") from exc

    wb = load_workbook(path, data_only=True, read_only=True)
    sheet_nodes: list[SectionNode] = []

    for sheet_idx, ws in enumerate(wb.worksheets, start=1):
        rows: list[tuple[int, str]] = []
        for i, row in enumerate(ws.iter_rows(values_only=True), start=1):
            vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if not vals:
                continue
            rows.append((i, " | ".join(vals)))

        sheet_summary = (
            " ".join(r[1] for r in rows) if rows else "[No extractable table/text rows in this sheet]"
        )

        # Chunk rows for tree children.
        chunk_size = 50
        children: list[SectionNode] = []
        for cstart in range(0, len(rows), chunk_size):
            chunk = rows[cstart : cstart + chunk_size]
            if not chunk:
                continue
            start_row = chunk[0][0]
            end_row = chunk[-1][0]
            children.append(
                SectionNode(
                    title=f"Rows {start_row}-{end_row}",
                    start_page=start_row,
                    end_page=end_row,
                    summary=" ".join(x[1] for x in chunk),
                )
            )

        sheet_nodes.append(
            SectionNode(
                title=f"Sheet: {ws.title}",
                start_page=sheet_idx,
                end_page=sheet_idx,
                summary=sheet_summary,
                children=children,
            )
        )

    stat = os.stat(path)
    metadata = {
        "file_name": os.path.basename(path),
        "file_path": path,
        "file_size_bytes": stat.st_size,
        "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "extension": os.path.splitext(path)[1].lower(),
        "indexing": {"index_source": "xlsx_sheets", "unit": "sheet/row"},
    }

    return DocumentIndex(
        title=sheet_nodes[0].title if sheet_nodes else None,
        total_pages=len(sheet_nodes),
        nodes=sheet_nodes,
        source=path,
        metadata=metadata,
    )

